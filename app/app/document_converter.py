# =============================================================================
# app/document_converter.py
# 
# Converts markdown content to .docx and .pdf files.
# Used by the /convert-document endpoint.
#
# Dependencies (add to Dockerfile):
#   pip install python-docx markdown xhtml2pdf
# =============================================================================

import io
import re
import markdown
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from xhtml2pdf import pisa
from bs4 import BeautifulSoup


def markdown_to_docx(md_content: str, title: str = "Document") -> bytes:
    """Convert markdown content to a .docx file. Returns bytes."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Convert markdown to HTML for parsing
    html = markdown.markdown(md_content, extensions=['tables', 'fenced_code', 'nl2br'])
    soup = BeautifulSoup(html, 'html.parser')
    
    for element in soup.children:
        _add_element_to_docx(doc, element)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _add_element_to_docx(doc: Document, element):
    """Recursively add HTML elements to a docx Document."""
    if isinstance(element, str):
        text = element.strip()
        if text:
            doc.add_paragraph(text)
        return
    
    tag = element.name
    
    if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
        level = int(tag[1])
        heading = doc.add_heading(element.get_text(), level=min(level, 4))
        # Style headings
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    
    elif tag == 'p':
        para = doc.add_paragraph()
        _add_inline_formatting(para, element)
    
    elif tag in ('ul', 'ol'):
        for i, li in enumerate(element.find_all('li', recursive=False)):
            para = doc.add_paragraph(style='List Bullet' if tag == 'ul' else 'List Number')
            _add_inline_formatting(para, li)
    
    elif tag == 'blockquote':
        text = element.get_text().strip()
        para = doc.add_paragraph(text)
        para.style = doc.styles['Intense Quote'] if 'Intense Quote' in [s.name for s in doc.styles] else doc.styles['Normal']
        para.paragraph_format.left_indent = Inches(0.5)
    
    elif tag == 'pre':
        code = element.get_text()
        para = doc.add_paragraph()
        run = para.add_run(code)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
        para.paragraph_format.left_indent = Inches(0.3)
    
    elif tag == 'table':
        _add_table_to_docx(doc, element)
    
    elif tag == 'hr':
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(12)
        run = para.add_run('─' * 60)
        run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
        run.font.size = Pt(8)
    
    else:
        # Recurse into unknown container elements
        for child in element.children:
            _add_element_to_docx(doc, child)


def _add_inline_formatting(para, element):
    """Handle inline formatting (bold, italic, code, links) within a paragraph."""
    for child in element.children:
        if isinstance(child, str):
            para.add_run(child)
        elif child.name == 'strong' or child.name == 'b':
            run = para.add_run(child.get_text())
            run.bold = True
        elif child.name == 'em' or child.name == 'i':
            run = para.add_run(child.get_text())
            run.italic = True
        elif child.name == 'code':
            run = para.add_run(child.get_text())
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xc7, 0x25, 0x4e)
        elif child.name == 'a':
            text = child.get_text()
            href = child.get('href', '')
            run = para.add_run(f"{text} ({href})" if href else text)
            run.font.color.rgb = RGBColor(0x06, 0x45, 0xAD)
            run.underline = True
        else:
            # Recurse
            _add_inline_formatting(para, child)


def _add_table_to_docx(doc: Document, table_element):
    """Convert an HTML table to a docx table."""
    rows_data = []
    
    # Get header rows
    thead = table_element.find('thead')
    if thead:
        for tr in thead.find_all('tr'):
            cells = [td.get_text().strip() for td in tr.find_all(['th', 'td'])]
            rows_data.append(cells)
    
    # Get body rows
    tbody = table_element.find('tbody') or table_element
    for tr in tbody.find_all('tr'):
        cells = [td.get_text().strip() for td in tr.find_all(['th', 'td'])]
        if cells:
            rows_data.append(cells)
    
    if not rows_data:
        return
    
    max_cols = max(len(row) for row in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=max_cols)
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            if j < max_cols:
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                # Bold header row
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True


def markdown_to_pdf(md_content: str, title: str = "Document") -> bytes:
    """Convert markdown content to a .pdf file. Returns bytes."""
    # Convert markdown to HTML
    html_body = markdown.markdown(md_content, extensions=['tables', 'fenced_code', 'nl2br'])
    
    # Wrap in a full HTML document with CSS for professional styling
    full_html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>{title}</title>
<style>
    @page {{
        size: letter;
        margin: 1in;
    }}
    body {{
        font-family: Helvetica, Arial, sans-serif;
        font-size: 11pt;
        line-height: 1.6;
        color: #333;
    }}
    h1 {{
        font-size: 22pt;
        color: #1a1a2e;
        border-bottom: 2px solid #e0e0e0;
        padding-bottom: 8px;
        margin-top: 24px;
    }}
    h2 {{
        font-size: 16pt;
        color: #1a1a2e;
        margin-top: 20px;
    }}
    h3 {{
        font-size: 13pt;
        color: #2d2d2d;
        margin-top: 16px;
    }}
    p {{
        margin: 8px 0;
    }}
    table {{
        width: 100%;
        border-collapse: collapse;
        margin: 12px 0;
        font-size: 10pt;
    }}
    th, td {{
        border: 1px solid #ddd;
        padding: 8px 10px;
        text-align: left;
    }}
    th {{
        background-color: #f5f5f5;
        font-weight: bold;
    }}
    tr:nth-child(even) {{
        background-color: #fafafa;
    }}
    blockquote {{
        border-left: 3px solid #ccc;
        margin: 12px 0;
        padding: 8px 16px;
        color: #555;
        background: #f9f9f9;
    }}
    code {{
        font-family: 'Courier New', monospace;
        font-size: 10pt;
        background: #f4f4f4;
        padding: 2px 4px;
        border-radius: 3px;
    }}
    pre {{
        background: #f4f4f4;
        padding: 12px;
        border-radius: 4px;
        overflow-x: auto;
        font-size: 9pt;
    }}
    pre code {{
        background: none;
        padding: 0;
    }}
    ul, ol {{
        margin: 8px 0;
        padding-left: 24px;
    }}
    li {{
        margin: 4px 0;
    }}
    hr {{
        border: none;
        border-top: 1px solid #e0e0e0;
        margin: 20px 0;
    }}
    strong {{
        color: #1a1a2e;
    }}
    a {{
        color: #0645AD;
        text-decoration: none;
    }}
</style>
</head>
<body>
{html_body}
</body>
</html>"""
    
    buffer = io.BytesIO()
    pisa_status = pisa.CreatePDF(full_html, dest=buffer)
    
    if pisa_status.err:
        raise ValueError(f"PDF generation failed with {pisa_status.err} errors")
    
    buffer.seek(0)
    return buffer.read()
